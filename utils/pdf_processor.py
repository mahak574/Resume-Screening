import pdfplumber
import PyPDF2
import docx
import io
import re
def extract_text_from_pdf(file_bytes):
    """Extracts text from a PDF file using pdfplumber, fallback to PyPDF2."""
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        # Fallback to PyPDF2 if pdfplumber fails
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e2:
            print(f"Failed to read PDF: {e2}")
    return text

def extract_text_from_docx(file_bytes):
    """Extracts text from a DOCX file using python-docx."""
    text = ""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        # Extract from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Failed to read DOCX: {e}")
        # Note: If this fails, it might be an older .doc file renamed to .docx.
    return text

def extract_text(file_upload):
    """Main extraction routing based on file extension."""
    file_name = file_upload.name.lower()
    file_bytes = file_upload.getvalue()
    
    if file_name.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif file_name.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    return ""

def parse_candidate_info(text: str) -> dict:
    """Parses text to extract name, email, phone, linkedin, and years of experience."""
    info = {
        "name": "Unknown",
        "email": "Not found",
        "phone": "Not found",
        "linkedin": "Not found",
        "years_experience": 0.0
    }
    
    if not text:
        return info
        
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if lines:
        for line in lines[:5]:
            clean_line = re.sub(r'[\+0-9\(\)-]+', '', line).strip() # Remove phone numbers
            clean_line = re.sub(r'[\w\.-]+@[\w\.-]+\.\w+', '', clean_line).strip() # Remove email
            
            # Collapse spaced out letters like "M a h a k  S i n g h" -> "Mahak Singh"
            # This is tricky, a simple replace of single spaces if surrounded by single letters
            if " " in clean_line and all(len(part) == 1 for part in clean_line.split()):
                clean_line = clean_line.replace(" ", "")
                
            lower_line = clean_line.lower()
            if 'curriculum vitae' in lower_line or 'resume' in lower_line or not clean_line:
                continue
                
            info["name"] = clean_line
            break
        
    # Email extraction
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    if email_match:
        info["email"] = email_match.group(0)
        
    # Phone extraction (simple heuristic for 10+ digits, optionally with separators)
    phone_match = re.search(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
    if phone_match:
        info["phone"] = phone_match.group(0)
        
    # LinkedIn extraction
    linkedin_match = re.search(r'(https?://)?(www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+/?', text)
    if linkedin_match:
        info["linkedin"] = linkedin_match.group(0)
        
    # Years of experience extraction
    exp_match = re.search(r'(\d+(?:\.\d+)?)\+?\s*(?:years? of experience|years? experience|yrs? of exp)', text, re.IGNORECASE)
    if exp_match:
        try:
            info["years_experience"] = float(exp_match.group(1))
        except ValueError:
            pass
            
    return info